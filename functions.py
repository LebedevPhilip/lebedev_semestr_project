import os

from docx import Document
from docx.shared import Pt


def change_font_size(document):
    """

    Функция  изменяет размер шрифта во всем документе

    """

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(14)


def change_font(document):
    """

    Изменяет шрифт во всем документе.

    """
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'


def change_line_spacing(document):
    """

    Изменяет межстрочный интервал во всем документе.

    """
    for paragraph in document.paragraphs:
        paragraph.line_spacing = Pt(1.5)